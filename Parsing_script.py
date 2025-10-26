#!/usr/bin/env python3
"""
Universal Document Parser Script (Finalized)
-------------------------------------------
Parses various common document formats (txt, csv, json, xml, pdf, docx, xlsx, html)
using standard and popular Python libraries. Aims for broad compatibility and
graceful error handling.

Features:
- Supports multiple formats: .txt, .csv, .json, .xml, .pdf, .docx, .xlsx, .xls, .html, .htm
- Detects format based on file extension.
- Conditional library imports with warnings if dependencies are missing.
- Basic text extraction for complex formats (PDF, DOCX, HTML, XML).
- Structured data loading for CSV, JSON, Excel.
- Error handling for file not found, decoding errors, and parsing issues.
- Returns parsed content in a format-appropriate Python object (string, list, dict).

Limitations:
- Encoding: Assumes UTF-8 by default for text files. May fail on files with
  different encodings. Consider using libraries like 'chardet' for auto-detection
  in more advanced scenarios.
- PDF Complexity: Text extraction via PyPDF2 works best on text-based PDFs.
  Scanned PDFs (images) require OCR (e.g., pytesseract). Complex layouts might
  result in imperfect text ordering.
- DOCX/XLSX Features: Primarily extracts text. Does not parse complex formatting,
  macros, images, or embedded objects. Pandas reads full Excel sheets into memory,
  which might be an issue for extremely large files.
- Structured Data: XML/HTML text extraction is basic. For specific data extraction,
  use targeted XPath (XML) or CSS selectors/tag navigation (HTML).
- CSV Dialects: Assumes standard comma-separated values. Use `csv.Sniffer` or
  pass specific `delimiter`, `quotechar` arguments for non-standard CSVs.

External Library Installation:
------------------------------
pip install PyPDF2 python-docx pandas openpyxl xlrd beautifulsoup4 lxml
"""

import os
import csv
import json
import xml.etree.ElementTree as ET
import sys
import logging
import pprint # For pretty printing JSON/Dicts

# --- Setup Logging ---
# Using logging is better than print for errors/info in reusable scripts
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')

# --- Check and Import External Libraries ---
# Suppress info messages during import checks
original_level = logging.getLogger().getEffectiveLevel()
logging.getLogger().setLevel(logging.WARNING)

# PyPDF2 for PDF
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False
    logging.warning("PyPDF2 not found. PDF parsing disabled. Install with: pip install PyPDF2")

# python-docx for DOCX
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.warning("python-docx not found. DOCX parsing disabled. Install with: pip install python-docx")

# pandas, openpyxl, xlrd for Excel
try:
    import pandas as pd
    import openpyxl # Needed by pandas for .xlsx
    import xlrd     # Needed by pandas for older .xls
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    logging.warning("pandas, openpyxl, or xlrd not found. Excel parsing disabled. Install with: pip install pandas openpyxl xlrd")

# BeautifulSoup4 and lxml for HTML
try:
    from bs4 import BeautifulSoup
    import lxml # Recommended parser for BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False
    logging.warning("beautifulsoup4 or lxml not found. HTML parsing disabled. Install with: pip install beautifulsoup4 lxml")

# Restore original logging level
logging.getLogger().setLevel(original_level)


# --- Parsing Functions ---

def parse_txt(filepath):
    """
    Parses a plain text file (.txt).
    Returns the file content as a single string, or None on error.
    """
    try:
        # Try UTF-8 first, then fall back to system default if it fails
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            logging.warning(f"UTF-8 decoding failed for {filepath}. Trying system default encoding.")
            with open(filepath, 'r', encoding=sys.getdefaultencoding(), errors='ignore') as f:
                content = f.read()
        logging.info(f"Successfully parsed TXT: {os.path.basename(filepath)}")
        # Optionally print snippet
        # print(f"--- Content Snippet of {os.path.basename(filepath)} ---")
        # print(content[:500] + ('...' if len(content) > 500 else ''))
        return content
    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
    except IOError as e:
        logging.error(f"IOError reading TXT file {filepath}: {e}")
    except Exception as e:
        logging.error(f"Unexpected error parsing TXT file {filepath}: {e}")
    return None

def parse_csv(filepath):
    """
    Parses a CSV file (.csv). Assumes comma delimiter and standard quoting.
    Returns a list of lists (header + data rows), or None on error.
    """
    try:
        rows = []
        with open(filepath, 'r', encoding='utf-8', newline='') as f:
            # Consider using csv.Sniffer().sniff(f.read(1024)) for dialect detection
            # f.seek(0) # Reset file pointer after sniffing
            reader = csv.reader(f)
            for row in reader:
                rows.append(row)
        logging.info(f"Successfully parsed CSV: {os.path.basename(filepath)} ({len(rows)} rows)")
        # Optionally print snippet
        # print(f"--- Header and first 5 data rows of {os.path.basename(filepath)} ---")
        # if rows:
        #     print("Header:", rows[0])
        #     for i, row in enumerate(rows[1:6], 1):
        #          print(f"Row {i}:", row)
        #     if len(rows) > 6:
        #          print(f"... and {len(rows)-6} more rows.")
        return rows
    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
    except csv.Error as e:
        logging.error(f"CSV parsing error in file {filepath}, line {reader.line_num}: {e}")
    except IOError as e:
        logging.error(f"IOError reading CSV file {filepath}: {e}")
    except Exception as e:
        logging.error(f"Unexpected error parsing CSV file {filepath}: {e}")
    return None

def parse_json(filepath):
    """
    Parses a JSON file (.json).
    Returns the parsed Python object (dict or list), or None on error.
    """
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        logging.info(f"Successfully parsed JSON: {os.path.basename(filepath)}")
        # Optionally print snippet
        # print(f"--- Parsed JSON data from {os.path.basename(filepath)} ---")
        # pprint.pprint(data)
        return data
    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
    except json.JSONDecodeError as e:
         logging.error(f"Invalid JSON format in {filepath}: {e}")
    except IOError as e:
        logging.error(f"IOError reading JSON file {filepath}: {e}")
    except Exception as e:
        logging.error(f"Unexpected error parsing JSON file {filepath}: {e}")
    return None

def parse_xml(filepath):
    """
    Parses an XML file (.xml) and extracts all text content.
    Returns the concatenated text as a single string, or None on error.
    """
    try:
        tree = ET.parse(filepath)
        root = tree.getroot()
        # Extract text content from all elements
        text_parts = [elem.text for elem in root.iter() if elem.text]
        text_content = ' '.join(part.strip() for part in text_parts if part.strip())
        logging.info(f"Successfully parsed XML and extracted text: {os.path.basename(filepath)}")
        # Optionally print snippet
        # print(f"--- Extracted Text from {os.path.basename(filepath)} ---")
        # print(text_content[:500] + ('...' if len(text_content) > 500 else ''))
        return text_content
    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
    except ET.ParseError as e:
        logging.error(f"XML parsing error in {filepath}: {e}")
    except IOError as e:
        logging.error(f"IOError reading XML file {filepath}: {e}")
    except Exception as e:
        logging.error(f"Unexpected error processing XML file {filepath}: {e}")
    return None

def parse_pdf(filepath):
    """
    Parses a PDF file (.pdf) and extracts text using PyPDF2.
    Returns the extracted text as a single string, or None on error or if library missing.
    """
    if not HAS_PYPDF2:
        logging.error("PyPDF2 library is required for PDF parsing but not found.")
        return None
    try:
        text_content = ""
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)
            logging.info(f"Reading {num_pages} pages from PDF: {os.path.basename(filepath)}...")
            for page_num in range(num_pages):
                try:
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n" # Add newline between pages
                except Exception as page_error:
                    # Log error for specific page but continue
                    logging.warning(f"Could not extract text from page {page_num + 1} in {filepath}: {page_error}")
        logging.info(f"Finished extracting text from PDF: {os.path.basename(filepath)}")
        # Optionally print snippet
        # print(f"--- Extracted Text from {os.path.basename(filepath)} ---")
        # print(text_content[:500] + ('...' if len(text_content) > 500 else ''))
        return text_content.strip()
    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
    except PyPDF2.errors.PdfReadError as e:
        logging.error(f"PyPDF2 error reading PDF {filepath}: {e}")
    except IOError as e:
        logging.error(f"IOError reading PDF file {filepath}: {e}")
    except Exception as e:
        logging.error(f"Unexpected error parsing PDF file {filepath}: {e}")
    return None

def parse_docx(filepath):
    """
    Parses a Word document (.docx) and extracts text from paragraphs.
    Returns the extracted text as a single string, or None on error or if library missing.
    """
    if not HAS_DOCX:
        logging.error("python-docx library is required for DOCX parsing but not found.")
        return None
    try:
        doc = docx.Document(filepath)
        text_content = "\n".join([para.text for para in doc.paragraphs if para.text])
        logging.info(f"Successfully parsed DOCX: {os.path.basename(filepath)}")
        # Optionally print snippet
        # print(f"--- Extracted Text from {os.path.basename(filepath)} ---")
        # print(text_content[:500] + ('...' if len(text_content) > 500 else ''))
        # TODO: Add table text extraction if needed:
        # for table in doc.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             text_content += cell.text + "\t" # Example separator
        #         text_content += "\n"
        return text_content
    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
    except Exception as e: # Catches errors from python-docx (e.g., corrupted file)
        logging.error(f"Error parsing DOCX file {filepath}: {e}")
    return None

def parse_excel(filepath):
    """
    Parses an Excel file (.xlsx, .xls) using pandas. Reads all sheets.
    Returns a dictionary mapping sheet names to pandas DataFrames, or None on error or if library missing.
    Note: Can consume significant memory for large files.
    """
    if not HAS_PANDAS:
        logging.error("pandas, openpyxl, xlrd libraries required for Excel parsing but not found.")
        return None
    try:
        # Determine the engine based on extension
        engine = 'openpyxl' if filepath.lower().endswith('.xlsx') else 'xlrd'
        excel_data = pd.read_excel(filepath, sheet_name=None, engine=engine)
        logging.info(f"Successfully parsed Excel: {os.path.basename(filepath)} ({len(excel_data)} sheets)")
        # Optionally print snippets
        # print(f"--- Content Snippets from {os.path.basename(filepath)} ---")
        # for sheet_name, df in excel_data.items():
        #     print(f"\n=== Sheet: {sheet_name} (first 5 rows) ===")
        #     print(df.head().to_string())
        return excel_data
    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
    except (KeyError, ValueError, ImportError) as e: # Catch pandas/engine errors
        logging.error(f"Error parsing Excel file {filepath} (check file/libraries): {e}")
    except MemoryError:
        logging.error(f"MemoryError parsing Excel file {filepath}. File may be too large for pandas.")
    except Exception as e:
        logging.error(f"Unexpected error parsing Excel file {filepath}: {e}")
    return None

def parse_html(filepath):
    """
    Parses an HTML file (.html, .htm) using BeautifulSoup and extracts text content.
    Returns the cleaned text as a single string, or None on error or if library missing.
    """
    if not HAS_BS4:
        logging.error("beautifulsoup4 and lxml libraries required for HTML parsing but not found.")
        return None
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            # Use 'lxml' for speed and robustness; 'html.parser' is built-in but slower
            soup = BeautifulSoup(f, 'lxml')

        # Remove script and style elements
        for element in soup(["script", "style"]):
            element.decompose() # Remove the tag from the tree

        # Get text, strip leading/trailing whitespace from lines, join non-empty lines
        text = soup.get_text(separator='\n', strip=True)

        logging.info(f"Successfully parsed HTML and extracted text: {os.path.basename(filepath)}")
        # Optionally print snippet
        # print(f"--- Extracted Text from {os.path.basename(filepath)} ---")
        # print(text[:500] + ('...' if len(text) > 500 else ''))
        return text
    except FileNotFoundError:
        logging.error(f"File not found: {filepath}")
    except IOError as e:
        logging.error(f"IOError reading HTML file {filepath}: {e}")
    except Exception as e: # Catch potential BeautifulSoup errors
        logging.error(f"Error parsing HTML file {filepath}: {e}")
    return None

# --- Main Parsing Dispatcher ---

def parse_document(filepath):
    """
    Parses a document based on its file extension. Logs errors.

    Args:
        filepath (str): The path to the document file.

    Returns:
        The parsed content in a format-specific type (str, list, dict, dict[str, pd.DataFrame]),
        or None if the file doesn't exist, the format is unsupported, or parsing fails.
    """
    if not isinstance(filepath, str) or not filepath:
        logging.error("Invalid filepath provided.")
        return None

    if not os.path.exists(filepath):
        logging.error(f"File does not exist: {filepath}")
        return None
    if not os.path.isfile(filepath):
         logging.error(f"Path is not a file: {filepath}")
         return None

    _, extension = os.path.splitext(filepath)
    extension = extension.lower()

    logging.info(f"Attempting to parse file: {filepath} (Detected type: {extension})")

    content = None
    if extension == '.txt':
        content = parse_txt(filepath)
    elif extension == '.csv':
        content = parse_csv(filepath)
    elif extension == '.json':
        content = parse_json(filepath)
    elif extension == '.xml':
        content = parse_xml(filepath)
    elif extension == '.pdf':
        content = parse_pdf(filepath)
    elif extension == '.docx':
        content = parse_docx(filepath)
    elif extension in ['.xlsx', '.xls']:
        content = parse_excel(filepath)
    elif extension in ['.html', '.htm']:
        content = parse_html(filepath)
    else:
        logging.warning(f"Unsupported file format '{extension}' for file: {filepath}")

    if content is not None:
         logging.info(f"Finished parsing: {os.path.basename(filepath)}")
    else:
         logging.error(f"Failed to parse: {os.path.basename(filepath)}")

    return content

# --- Example Usage ---

if __name__ == "__main__":
    logging.info("Starting document parsing example...")

    # Create dummy files for testing (overwrite if they exist)
    dummy_files_created = []
    try:
        logging.info("Creating dummy files...")
        # TXT
        with open("sample.txt", "w", encoding='utf-8') as f: f.write("This is line one.\nThis is line two.")
        dummy_files_created.append("sample.txt")
        # CSV
        with open("sample.csv", "w", newline="", encoding='utf-8') as f:
            writer = csv.writer(f); writer.writerow(["ColA", "ColB"]); writer.writerow(["Data1", 10]); writer.writerow(["Data2", 20])
        dummy_files_created.append("sample.csv")
        # JSON
        with open("sample.json", "w", encoding='utf-8') as f: json.dump({"key": "value", "list": [1, None, True]}, f)
        dummy_files_created.append("sample.json")
        # XML
        with open("sample.xml", "w", encoding='utf-8') as f: f.write('<data><item name="A">1</item><item name="B">2</item></data>')
        dummy_files_created.append("sample.xml")
        # HTML
        with open("sample.html", "w", encoding='utf-8') as f: f.write('<!DOCTYPE html><html><body><p>Hello</p><p>World</p></body></html>')
        dummy_files_created.append("sample.html")
        logging.info(f"Dummy files created: {', '.join(dummy_files_created)}")
    except Exception as e:
        logging.error(f"Error creating dummy files: {e}")

    # --- Files to Parse (Update paths for PDF, DOCX, XLSX) ---
    files_to_parse = [
        "sample.txt",
        "sample.csv",
        "sample.json",
        "sample.xml",
        "sample.html",
        # --- ADD PATHS TO YOUR REAL FILES BELOW ---
        # "path/to/your/document.pdf",
        # "path/to/your/document.docx",
        # "path/to/your/spreadsheet.xlsx",
        # ------------------------------------------
        "non_existent_file.xyz" # Test file not found
    ]

    results = {}
    print("\n" + "=" * 70)
    logging.info("Starting batch parsing...")
    print("=" * 70)

    for file in files_to_parse:
        parsed_data = parse_document(file)
        results[file] = parsed_data
        print("-" * 70) # Separator between file outputs

    logging.info("Batch parsing finished.")
    print("=" * 70)

    # --- Accessing Results Example ---
    logging.info("Example: Accessing parsed data from results dictionary...")
    txt_content = results.get("sample.txt")
    if txt_content:
        logging.info(f"Content of sample.txt: '{txt_content[:30]}...'")

    csv_data = results.get("sample.csv")
    if csv_data and len(csv_data) > 1:
        logging.info(f"First data row of sample.csv: {csv_data[1]}")

    json_data = results.get("sample.json")
    if json_data:
        logging.info(f"Value of 'key' in sample.json: {json_data.get('key')}")

    # --- Clean up dummy files ---
    logging.info("Cleaning up dummy files...")
    cleaned_count = 0
    for dummy in dummy_files_created:
         try:
             if os.path.exists(dummy):
                 os.remove(dummy)
                 cleaned_count += 1
         except Exception as e:
             logging.warning(f"Could not remove dummy file {dummy}: {e}")
    logging.info(f"Removed {cleaned_count} dummy files.")

    logging.info("Script finished.")
